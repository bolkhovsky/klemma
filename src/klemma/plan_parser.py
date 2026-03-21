"""Parser for project plan / outline files.

Extracts structured data from academic project outlines:
title, description (актуальность), research objectives, scientific results,
and full chapter structure.

Supported formats:
- .docx — full structured parsing from dissertation plan-prospect (python-docx required)
- .md / .txt — text-based structure extraction from headings and numbered lists
- Raw text — same extraction via parse_text()

Designed as a pure module — no CLI dependencies. Works with file paths
or BytesIO for SaaS use.
"""

import re
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Optional, Union


@dataclass
class Section:
    """A section/subsection in the dissertation structure."""

    number: str          # e.g. "1.1", "2.3.1"
    title: str
    level: int           # 1=chapter, 2=section, 3=subsection
    children: list["Section"] = field(default_factory=list)


@dataclass
class Chapter:
    """A chapter in the dissertation structure."""

    number: int          # 1, 2, 3, 4
    title: str
    sections: list[Section] = field(default_factory=list)


@dataclass
class ScientificResult:
    """A scientific result (НР) from the plan."""

    number: int          # 1, 2, ...
    title: str           # short title from Heading 2
    description: str     # full description from body text


@dataclass
class PlanData:
    """Structured data extracted from a dissertation plan-prospect."""

    title: str = ""
    description: str = ""          # актуальность
    research_object: str = ""      # объект исследования
    research_subject: str = ""     # предмет исследования
    research_goal: str = ""        # цель исследования
    tasks: list[str] = field(default_factory=list)
    results: list[ScientificResult] = field(default_factory=list)
    chapters: list[Chapter] = field(default_factory=list)
    language: str = "ru"


# --- Heading markers for section detection ---
_HEADING_MARKERS = {
    "актуальность": "description",
    "объект исследования": "research_object",
    "предмет исследования": "research_subject",
    "цель исследования": "research_goal",
    "задачи исследования": "tasks",
    "выносимые на защиту": "results",
    "содержание диссертации": "chapters",
}

# Regex for chapter lines: "Глава 1. Title"
_RE_CHAPTER = re.compile(r"^Глава\s+(\d+)\.\s*(.+)")

# Regex for section/subsection lines: "- 1.1. Title" or "1.1. Title"
_RE_SECTION = re.compile(r"^-?\s*(\d+(?:\.\d+)+)\.?\s*(.+)")

# Regex for "Выводы по главе N"
_RE_CONCLUSIONS = re.compile(r"^Выводы по главе\s+\d+", re.IGNORECASE)

# Regex for НР heading: "НР 1. Title"
_RE_NR = re.compile(r"^НР\s+(\d+)\.\s*(.+)")

# Structural lines to skip in chapter parsing
_STRUCTURAL_LINES = {"ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ЛИТЕРАТУРЫ", "ПРИЛОЖЕНИЯ"}


def parse(source: Union[str, Path, BytesIO]) -> PlanData:
    """Parse a dissertation plan-prospect .docx file.

    Args:
        source: file path (str or Path) or BytesIO with .docx content.

    Returns:
        PlanData with extracted fields.

    Raises:
        ImportError: if python-docx is not installed.
        FileNotFoundError: if path doesn't exist.
        ValueError: if file cannot be parsed.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required to parse .docx files. "
            "Install it with: pip install python-docx"
        )

    if isinstance(source, (str, Path)):
        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"Plan file not found: {path}")
        doc = Document(str(path))
    else:
        doc = Document(source)

    paragraphs = doc.paragraphs
    data = PlanData()

    # --- Extract title ---
    data.title = _extract_title(paragraphs)

    # --- Split paragraphs by Heading 1 sections ---
    sections = _split_by_headings(paragraphs)

    for heading_text, body_paragraphs in sections:
        heading_lower = heading_text.lower()

        # Match heading to known markers
        matched_field = None
        for marker, field_name in _HEADING_MARKERS.items():
            if marker in heading_lower:
                matched_field = field_name
                break

        if matched_field == "description":
            data.description = _join_paragraphs(body_paragraphs)
        elif matched_field == "research_object":
            data.research_object = _join_paragraphs(body_paragraphs)
        elif matched_field == "research_subject":
            data.research_subject = _join_paragraphs(body_paragraphs)
        elif matched_field == "research_goal":
            data.research_goal = _join_paragraphs(body_paragraphs)
        elif matched_field == "tasks":
            data.tasks = _extract_tasks(body_paragraphs)
        elif matched_field == "results":
            data.results = _extract_results(body_paragraphs)
        elif matched_field == "chapters":
            data.chapters = _parse_chapters(body_paragraphs)

    return data


# --- Alternative chapter heading patterns for text/markdown ---
_RE_CHAPTER_EN = re.compile(r"^(?:Chapter|Раздел)\s+(\d+)[\.:\s]\s*(.+)", re.IGNORECASE)
_RE_CHAPTER_PLAIN = re.compile(r"^(\d+)[\.:\s]\s+(.+)")


def parse_text(text: str) -> PlanData:
    """Parse outline structure from plain text or markdown.

    Extracts chapter/section structure from headings and numbered patterns.
    Supports Russian (Глава N.) and English (Chapter N.) formats,
    as well as plain numbered lists (1. Title, 1.1. Section).

    Args:
        text: outline text (plain or markdown).

    Returns:
        PlanData with extracted fields (may have empty chapters if
        no structure detected — the raw text is still useful as context).
    """
    data = PlanData()
    lines = text.strip().split("\n")

    if not lines:
        return data

    current_chapter: Optional[Chapter] = None
    title_found = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Remove markdown heading markers for pattern matching
        clean = stripped.lstrip("#").strip()

        # First significant heading → title
        if not title_found and stripped.startswith("#") and not stripped.startswith("##"):
            data.title = clean
            title_found = True
            continue

        # Try chapter patterns: "Глава N. Title"
        ch_match = _RE_CHAPTER.match(clean)
        if not ch_match:
            ch_match = _RE_CHAPTER_EN.match(clean)
        if not ch_match:
            # "N. Title" (single number, not "N.N. Title")
            m = _RE_CHAPTER_PLAIN.match(clean)
            if m and "." not in m.group(1):
                ch_match = m

        if ch_match:
            num = int(ch_match.group(1))
            title = ch_match.group(2).strip()
            if title.upper() in _STRUCTURAL_LINES:
                continue
            current_chapter = Chapter(number=num, title=title)
            data.chapters.append(current_chapter)
            continue

        # Section: "N.N. Title" or "- N.N. Title"
        sec_match = _RE_SECTION.match(clean)
        if sec_match and current_chapter is not None:
            sec_num = sec_match.group(1)
            sec_title = sec_match.group(2).strip()
            if _RE_CONCLUSIONS.match(clean):
                continue
            level = sec_num.count(".") + 1
            section = Section(number=sec_num, title=sec_title, level=level)
            current_chapter.sections.append(section)
            continue

        # Scientific results: "НР N. Title"
        nr_match = _RE_NR.match(clean)
        if nr_match:
            data.results.append(ScientificResult(
                number=int(nr_match.group(1)),
                title=nr_match.group(2).strip(),
                description="",
            ))

    # Fallback title: first non-empty line
    if not data.title:
        for line in lines:
            s = line.strip().lstrip("#").strip()
            if s:
                data.title = s
                break

    return data


def parse_file(source: Union[str, Path]) -> PlanData:
    """Parse plan/outline from file, auto-detecting format by extension.

    Supported formats:
    - .docx — full structured parsing (python-docx required)
    - .md, .txt, .text, .markdown — text-based structure extraction

    Args:
        source: file path.

    Returns:
        PlanData with extracted fields.

    Raises:
        FileNotFoundError: if file doesn't exist.
        ValueError: if format is not supported.
        ImportError: if .docx and python-docx not installed.
    """
    path = Path(source)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    ext = path.suffix.lower()
    if ext == ".docx":
        return parse(path)
    elif ext in (".md", ".txt", ".text", ".markdown"):
        text = path.read_text(encoding="utf-8")
        return parse_text(text)
    else:
        raise ValueError(
            f"Unsupported format: {ext}. Use .docx, .md, or .txt"
        )


def _extract_title(paragraphs) -> str:
    """Extract dissertation title from 'на тему «...»' paragraph."""
    for p in paragraphs[:10]:
        text = p.text.strip()
        # Look for «...» quoted title
        m = re.search(r"[«\"](.*?)[»\"]", text)
        if m and len(m.group(1)) > 20:
            return m.group(1)
        # Look for "на тему" prefix with title continuing
        if "на тему" in text.lower():
            # Title might be after «
            m = re.search(r"[«\"](.*?)$", text)
            if m:
                title = m.group(1).rstrip("»\"")
                # Check next paragraphs for continuation
                idx = list(paragraphs).index(p)
                for next_p in paragraphs[idx + 1 : idx + 3]:
                    next_text = next_p.text.strip()
                    if next_text and not next_text.startswith("на соискание"):
                        title += " " + next_text.rstrip("»\"")
                    else:
                        break
                return title
    return ""


def _split_by_headings(paragraphs) -> list[tuple[str, list]]:
    """Split paragraphs into (heading_text, [body_paragraphs]) groups by Heading 1."""
    sections: list[tuple[str, list]] = []
    current_heading = ""
    current_body: list = []

    for p in paragraphs:
        if p.style and p.style.name == "Heading 1":
            if current_heading:
                sections.append((current_heading, current_body))
            current_heading = p.text.strip()
            current_body = []
        elif current_heading:
            current_body.append(p)

    if current_heading:
        sections.append((current_heading, current_body))

    return sections


def _join_paragraphs(paragraphs) -> str:
    """Join paragraph texts into a single string."""
    parts = [p.text.strip() for p in paragraphs if p.text.strip()]
    return " ".join(parts)


def _extract_tasks(paragraphs) -> list[str]:
    """Extract numbered research tasks from paragraphs."""
    tasks = []
    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # Tasks are typically numbered list items
        # Clean leading markers like "1.", "1)", "-"
        cleaned = re.sub(r"^\d+[.)]\s*", "", text)
        cleaned = re.sub(r"^[-–—]\s*", "", cleaned)
        if cleaned and len(cleaned) > 10:
            tasks.append(cleaned)
    return tasks


def _extract_results(paragraphs) -> list[ScientificResult]:
    """Extract scientific results (НР) from Heading 2 + body pairs."""
    results: list[ScientificResult] = []
    current_nr: Optional[ScientificResult] = None

    for p in paragraphs:
        if p.style and p.style.name == "Heading 2":
            # Save previous
            if current_nr:
                results.append(current_nr)
            # Parse "НР N. Title"
            m = _RE_NR.match(p.text.strip())
            if m:
                current_nr = ScientificResult(
                    number=int(m.group(1)),
                    title=m.group(2).strip(),
                    description="",
                )
            else:
                current_nr = ScientificResult(
                    number=len(results) + 1,
                    title=p.text.strip(),
                    description="",
                )
        elif current_nr:
            text = p.text.strip()
            if text:
                if current_nr.description:
                    current_nr.description += " " + text
                else:
                    current_nr.description = text

    if current_nr:
        results.append(current_nr)

    return results


def _parse_chapters(paragraphs) -> list[Chapter]:
    """Parse chapter structure from 'Содержание диссертации' section.

    Handles multi-line entries where a title wraps across paragraphs.
    Recognizes: 'Глава N. Title', '- N.N. Title', '- N.N.N. Title',
    'Выводы по главе N', structural markers (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, etc.)
    """
    # First, join continuation lines into complete entries
    lines = _join_continuation_lines(paragraphs)

    chapters: list[Chapter] = []
    current_chapter: Optional[Chapter] = None

    for line in lines:
        # Skip structural markers
        if line.upper() in _STRUCTURAL_LINES:
            continue
        if _RE_CONCLUSIONS.match(line):
            continue
        # Skip appendices
        if line.startswith("Приложение") or line.startswith("- Приложение"):
            continue

        # Chapter: "Глава N. Title"
        m_ch = _RE_CHAPTER.match(line)
        if m_ch:
            current_chapter = Chapter(
                number=int(m_ch.group(1)),
                title=m_ch.group(2).strip(),
            )
            chapters.append(current_chapter)
            continue

        # Section: "- N.N. Title" or "- N.N.N. Title"
        m_sec = _RE_SECTION.match(line)
        if m_sec and current_chapter is not None:
            num = m_sec.group(1)
            title = m_sec.group(2).strip()
            parts = num.split(".")
            level = len(parts)

            section = Section(number=num, title=title, level=level)

            if level == 2:
                current_chapter.sections.append(section)
            elif level >= 3 and current_chapter.sections:
                # Find parent section
                parent_num = ".".join(parts[:2])
                parent = None
                for s in current_chapter.sections:
                    if s.number == parent_num:
                        parent = s
                        break
                if parent:
                    parent.children.append(section)
                else:
                    current_chapter.sections.append(section)

    return chapters


def _join_continuation_lines(paragraphs) -> list[str]:
    """Join paragraphs that are continuations of the previous line.

    A paragraph is a continuation if it doesn't start with a recognized
    pattern (Глава, - N.N., Выводы, structural marker, Приложение).
    """
    lines: list[str] = []

    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # Is this a new entry or continuation?
        is_new = (
            _RE_CHAPTER.match(text) is not None
            or _RE_SECTION.match(text) is not None
            or _RE_CONCLUSIONS.match(text) is not None
            or text.upper() in _STRUCTURAL_LINES
            or text.startswith("Приложение")
            or text.startswith("- Приложение")
        )

        if is_new or not lines:
            lines.append(text)
        else:
            # Continuation of previous line
            lines[-1] += " " + text

    return lines


def to_klemma_md(data: PlanData) -> str:
    """Convert PlanData to KLEMMA.md content."""
    lines = [
        "# Project Context\n",
        "<!-- Extracted from dissertation plan-prospect. -->\n",
        f'Topic: "{data.title}"\n',
    ]

    if data.description:
        desc = data.description[:500]
        if len(data.description) > 500:
            desc += "..."
        lines.append(f"Description: {desc}\n")

    if data.research_goal:
        lines.append(f"Goal: {data.research_goal}\n")

    if data.results:
        lines.append("Scientific Results:")
        for nr in data.results:
            lines.append(f"- NR{nr.number}: {nr.title}")
        lines.append("")

    if data.tasks:
        lines.append("Research Tasks:")
        for i, task in enumerate(data.tasks, 1):
            task_short = task[:120] + "..." if len(task) > 120 else task
            lines.append(f"- T{i}: {task_short}")
        lines.append("")

    if data.chapters:
        lines.append("Structure:")
        for ch in data.chapters:
            lines.append(f"- Chapter {ch.number}: {ch.title}")
            for sec in ch.sections:
                lines.append(f"  - {sec.number}. {sec.title}")
        lines.append("")

    return "\n".join(lines)
